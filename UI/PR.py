import sys
import os
from pathlib import Path
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QPushButton, QLineEdit, QLabel, QFileDialog, QProgressBar,
    QMessageBox, QRadioButton, QButtonGroup, QTextEdit, QSplitter,
    QTabWidget, QSpinBox, QGroupBox, QFormLayout
)
from PyQt5.QtCore import Qt, QThread, pyqtSignal
import docx
from docx2pdf import convert
import pdfplumber
import fitz  # PyMuPDF
from tqdm import tqdm


# -------------------------- 后台线程（转换+阅读） --------------------------
class ConverterThread(QThread):
    progress_updated = pyqtSignal(int)
    task_finished = pyqtSignal(str)
    task_failed = pyqtSignal(str)

    def __init__(self, input_path, output_path, conversion_type):
        super().__init__()
        self.input_path = input_path
        self.output_path = output_path
        self.conversion_type = conversion_type

    def run(self):
        try:
            if self.conversion_type == "word2pdf":
                self._convert_word_to_pdf()
            elif self.conversion_type == "pdf2word":
                self._convert_pdf_to_word()
            self.task_finished.emit(self.output_path)
        except Exception as e:
            self.task_failed.emit(str(e))

    def _convert_word_to_pdf(self):
        self.progress_updated.emit(20)
        convert(str(self.input_path), str(self.output_path))
        self.progress_updated.emit(100)

    def _convert_pdf_to_word(self):
        doc = docx.Document()
        with pdfplumber.open(self.input_path) as pdf:
            total_pages = len(pdf.pages)
            for i, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text:
                    doc.add_paragraph(text)
                self.progress_updated.emit(int((i + 1) / total_pages * 100))
        doc.save(str(self.output_path))


class DocumentReaderThread(QThread):
    """后台读取文档内容线程"""
    content_loaded = pyqtSignal(str, int)  # 内容、总页数
    page_loaded = pyqtSignal(str)  # 单页内容
    load_failed = pyqtSignal(str)

    def __init__(self, file_path, file_type):
        super().__init__()
        self.file_path = file_path
        self.file_type = file_type  # "pdf" 或 "word"

    def run(self):
        try:
            if self.file_type == "pdf":
                self._read_pdf()
            elif self.file_type == "word":
                self._read_word()
        except Exception as e:
            self.load_failed.emit(str(e))

    def _read_pdf(self):
        """读取PDF全部内容和页数"""
        content = ""
        total_pages = 0
        with fitz.open(self.file_path) as doc:
            total_pages = doc.page_count
            for page in doc:
                content += page.get_text() + "\n\n"
        self.content_loaded.emit(content, total_pages)

    def _read_word(self):
        """读取Word全部内容"""
        doc = docx.Document(self.file_path)
        content = "\n".join([p.text for p in doc.paragraphs])
        self.content_loaded.emit(content, 1)  # Word文档按单页处理


# -------------------------- 主界面窗口（含阅读功能） --------------------------
class DocumentToolUI(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("PDF/Word 全能处理工具")
        self.setGeometry(200, 200, 900, 600)
        self.current_file = None
        self.current_file_type = None
        self.all_content = ""
        self.total_pages = 1
        self.init_ui()

    def init_ui(self):
        # 主部件和布局
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        main_layout = QVBoxLayout(central_widget)
        main_layout.setSpacing(10)
        main_layout.setContentsMargins(10, 10, 10, 10)

        # 1. 功能选项卡
        self.tab_widget = QTabWidget()
        self.conversion_tab = QWidget()
        self.reader_tab = QWidget()

        self.tab_widget.addTab(self.conversion_tab, "格式转换")
        self.tab_widget.addTab(self.reader_tab, "文档阅读")

        main_layout.addWidget(self.tab_widget)

        # -------------------------- 转换功能界面 --------------------------
        self.init_conversion_ui()

        # -------------------------- 阅读功能界面 --------------------------
        self.init_reader_ui()

    def init_conversion_ui(self):
        """初始化格式转换界面（复用原有逻辑）"""
        layout = QVBoxLayout(self.conversion_tab)
        layout.setSpacing(15)
        layout.setContentsMargins(30, 30, 30, 30)

        # 转换类型选择
        type_layout = QHBoxLayout()
        self.word2pdf_radio = QRadioButton("Word → PDF")
        self.pdf2word_radio = QRadioButton("PDF → Word")
        self.word2pdf_radio.setChecked(True)
        type_group = QButtonGroup()
        type_group.addButton(self.word2pdf_radio)
        type_group.addButton(self.pdf2word_radio)

        type_layout.addWidget(self.word2pdf_radio)
        type_layout.addWidget(self.pdf2word_radio)
        layout.addLayout(type_layout)

        # 文件选择区
        input_layout = QHBoxLayout()
        self.input_path_edit = QLineEdit()
        self.input_path_edit.setPlaceholderText("请选择输入文件...")
        self.browse_input_btn = QPushButton("浏览")
        self.browse_input_btn.clicked.connect(self.browse_input_file)

        input_layout.addWidget(self.input_path_edit)
        input_layout.addWidget(self.browse_input_btn)
        layout.addLayout(input_layout)

        # 输出路径区
        output_layout = QHBoxLayout()
        self.output_path_edit = QLineEdit()
        self.output_path_edit.setPlaceholderText("输出文件路径...")
        self.browse_output_btn = QPushButton("浏览")
        self.browse_output_btn.clicked.connect(self.browse_output_path)

        output_layout.addWidget(self.output_path_edit)
        output_layout.addWidget(self.browse_output_btn)
        layout.addLayout(output_layout)

        # 进度条
        self.progress_bar = QProgressBar()
        self.progress_bar.setValue(0)
        layout.addWidget(self.progress_bar)

        # 操作按钮
        btn_layout = QHBoxLayout()
        self.convert_btn = QPushButton("开始转换")
        self.convert_btn.clicked.connect(self.start_conversion)
        self.convert_btn.setStyleSheet("background-color: #4CAF50; color: white; padding: 8px;")

        self.clear_btn = QPushButton("清空")
        self.clear_btn.clicked.connect(self.clear_fields)

        btn_layout.addWidget(self.convert_btn)
        btn_layout.addWidget(self.clear_btn)
        layout.addLayout(btn_layout)

    def init_reader_ui(self):
        """初始化文档阅读界面"""
        layout = QVBoxLayout(self.reader_tab)
        layout.setSpacing(10)

        # 顶部控制区
        control_layout = QHBoxLayout()

        # 文件选择
        self.reader_file_edit = QLineEdit()
        self.reader_file_edit.setPlaceholderText("请选择要阅读的文档...")
        self.browse_reader_btn = QPushButton("打开文档")
        self.browse_reader_btn.clicked.connect(self.browse_reader_file)

        # 分页控制
        self.page_spin = QSpinBox()
        self.page_spin.setMinimum(1)
        self.page_spin.valueChanged.connect(self.jump_to_page)

        # 搜索功能
        self.search_edit = QLineEdit()
        self.search_edit.setPlaceholderText("搜索文本...")
        self.search_btn = QPushButton("搜索")
        self.search_btn.clicked.connect(self.search_text)

        control_layout.addWidget(self.reader_file_edit)
        control_layout.addWidget(self.browse_reader_btn)
        control_layout.addWidget(QLabel("页码:"))
        control_layout.addWidget(self.page_spin)
        control_layout.addWidget(self.search_edit)
        control_layout.addWidget(self.search_btn)
        layout.addLayout(control_layout)

        # 内容显示区
        self.content_text = QTextEdit()
        self.content_text.setReadOnly(True)
        self.content_text.setStyleSheet("font-size: 12pt; line-height: 1.5;")
        layout.addWidget(self.content_text)

    # -------------------------- 转换功能逻辑 --------------------------
    def browse_input_file(self):
        file_filter = "Word文件 (*.docx *.doc)" if self.word2pdf_radio.isChecked() else "PDF文件 (*.pdf)"
        file_path, _ = QFileDialog.getOpenFileName(self, "选择输入文件", "", file_filter)
        if file_path:
            self.input_path_edit.setText(file_path)
            self._auto_set_output_path(file_path)

    def browse_output_path(self):
        file_filter = "PDF文件 (*.pdf)" if self.word2pdf_radio.isChecked() else "Word文件 (*.docx)"
        file_path, _ = QFileDialog.getSaveFileName(self, "选择输出路径", "", file_filter)
        if file_path:
            self.output_path_edit.setText(file_path)

    def _auto_set_output_path(self, input_path):
        input_path = Path(input_path)
        suffix = ".pdf" if self.word2pdf_radio.isChecked() else ".docx"
        output_path = input_path.with_suffix(suffix)
        self.output_path_edit.setText(str(output_path))

    def start_conversion(self):
        input_path = self.input_path_edit.text().strip()
        output_path = self.output_path_edit.text().strip()

        if not input_path or not output_path:
            QMessageBox.warning(self, "警告", "请选择输入和输出文件路径")
            return
        if not Path(input_path).exists():
            QMessageBox.warning(self, "错误", "输入文件不存在")
            return

        self.convert_btn.setEnabled(False)
        self.progress_bar.setValue(0)

        conversion_type = "word2pdf" if self.word2pdf_radio.isChecked() else "pdf2word"
        self.thread = ConverterThread(input_path, output_path, conversion_type)
        self.thread.progress_updated.connect(self.progress_bar.setValue)
        self.thread.task_finished.connect(self.on_conversion_finished)
        self.thread.task_failed.connect(self.on_conversion_failed)
        self.thread.start()

    def on_conversion_finished(self, output_path):
        self.convert_btn.setEnabled(True)
        QMessageBox.information(self, "完成", f"转换成功！\n文件已保存到:\n{output_path}")
        self.progress_bar.setValue(0)

    def on_conversion_failed(self, error_msg):
        self.convert_btn.setEnabled(True)
        QMessageBox.critical(self, "错误", f"转换失败:\n{error_msg}")
        self.progress_bar.setValue(0)

    def clear_fields(self):
        self.input_path_edit.clear()
        self.output_path_edit.clear()
        self.progress_bar.setValue(0)

    # -------------------------- 阅读功能逻辑 --------------------------
    def browse_reader_file(self):
        file_filter = "所有文档 (*.pdf *.docx *.doc)"
        file_path, _ = QFileDialog.getOpenFileName(self, "选择要阅读的文档", "", file_filter)
        if file_path:
            self.reader_file_edit.setText(file_path)
            self.current_file = file_path
            self.current_file_type = "pdf" if file_path.endswith((".pdf", ".PDF")) else "word"
            self.load_document()

    def load_document(self):
        """后台加载文档内容"""
        self.content_text.clear()
        self.page_spin.setValue(1)

        self.reader_thread = DocumentReaderThread(self.current_file, self.current_file_type)
        self.reader_thread.content_loaded.connect(self.on_content_loaded)
        self.reader_thread.load_failed.connect(self.on_reader_failed)
        self.reader_thread.start()

    def on_content_loaded(self, content, total_pages):
        """文档加载完成处理"""
        self.all_content = content
        self.total_pages = total_pages
        self.page_spin.setMaximum(total_pages)
        self.jump_to_page(1)  # 显示第一页

    def on_reader_failed(self, error_msg):
        QMessageBox.critical(self, "错误", f"文档加载失败:\n{error_msg}")

    def jump_to_page(self, page_num):
        """跳转到指定页码"""
        if not self.all_content:
            return

        if self.current_file_type == "pdf":
            # PDF按分页显示
            with fitz.open(self.current_file) as doc:
                if 1 <= page_num <= doc.page_count:
                    page = doc[page_num - 1]
                    self.content_text.setPlainText(page.get_text())
        else:
            # Word文档显示全部内容
            self.content_text.setPlainText(self.all_content)

    def search_text(self):
        """搜索文档内容"""
        search_text = self.search_edit.text().strip()
        if not search_text or not self.all_content:
            return

        # 高亮显示搜索结果
        self.content_text.find(search_text, QTextEdit.FindCaseSensitively)


# -------------------------- 运行应用 --------------------------
if __name__ == "__main__":
    os.environ['QT_QPA_PLATFORM_PLUGIN_PATH'] = ""

    app = QApplication(sys.argv)
    window = DocumentToolUI()
    window.show()
    sys.exit(app.exec_())